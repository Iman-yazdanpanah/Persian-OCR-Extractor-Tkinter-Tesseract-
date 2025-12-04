import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageGrab
import pytesseract
from docx import Document
import os
import sys


# new imports for proper display
import arabic_reshaper
from bidi.algorithm import get_display

current_image_name = None   # To store default filename

import pytesseract
import os


pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"

import subprocess
print(subprocess.check_output([
    r"C:\Program Files\Tesseract-OCR\tesseract.exe",
    "--list-langs"
]).decode())


# ---------------------------
# OCR FUNCTION (FROM IMAGE FILE)
# ---------------------------
def run_ocr():
    """Called when user selects image file."""
    global current_image_name, last_raw_text

    file_path = filedialog.askopenfilename(
        title="انتخاب تصویر",
        filetypes=[("Image files", "*.jpg *.jpeg *.png *.bmp *.tiff")]
    )
    if not file_path:
        return

    try:
        img = Image.open(file_path)
        current_image_name = os.path.splitext(os.path.basename(file_path))[0]
        run_ocr_on_image(img)   # this will set last_raw_text and update UI

    except Exception as e:
        messagebox.showerror("Error", str(e))


# ---------------------------
# OCR FROM CLIPBOARD
# ---------------------------
def paste_from_clipboard():
    import win32clipboard
    import win32con
    import struct
    import io

    global current_image_name

    try:
        data = ImageGrab.grabclipboard()

        if isinstance(data, Image.Image):
            current_image_name = "ocr_output"
            run_ocr_on_image(data)
            return

        if isinstance(data, list) and len(data) > 0:
            img = Image.open(data[0])
            current_image_name = os.path.splitext(os.path.basename(data[0]))[0]
            run_ocr_on_image(img)
            return

        win32clipboard.OpenClipboard()
        if win32clipboard.IsClipboardFormatAvailable(win32con.CF_DIB):
            dib_data = win32clipboard.GetClipboardData(win32con.CF_DIB)
            win32clipboard.CloseClipboard()

            if dib_data is None:
                raise Exception("Clipboard DIB is empty")

            bmp_header = b"BM"

            dib_size = len(dib_data)
            file_size = dib_size + 14  

            pixel_offset = 14 + struct.unpack_from("<I", dib_data, 0)[0]

            bmp_header += struct.pack("<IHHI", file_size, 0, 0, pixel_offset)
            bmp_bytes = bmp_header + dib_data

            img = Image.open(io.BytesIO(bmp_bytes))

            current_image_name = "ocr_output"
            run_ocr_on_image(img)
            return

        win32clipboard.CloseClipboard()
        messagebox.showwarning("Warning", "Clipboard does not contain an image")

    except Exception as e:
        try:
            win32clipboard.CloseClipboard()
        except:
            pass
        messagebox.showerror("Error", f"Clipboard error:\n{e}")




def run_ocr_on_image(img):
    global last_raw_text, current_image_name

    try:
        img = img.convert("RGB")

        # Path to tessdata inside PyInstaller EXE
        tess_config = "--oem 1 --psm 6"


        # Use Persian + English languages
        raw_text = pytesseract.image_to_string(
            img,
            lang="fas+eng",
            config=tess_config
        )

        last_raw_text = raw_text

        # Show raw text directly in Tkinter
        set_text_rtl(raw_text)

    except Exception as e:
        messagebox.showerror("Error", f"OCR error:\n{e}")


# ---------------------------
# RTL TEXT HANDLING
# ---------------------------
def set_text_rtl(text):
    text_box.delete(1.0, tk.END)
    text_box.insert(tk.END, text)
    text_box.tag_add("right", "1.0", "end")
    text_box.tag_configure("right", justify='right')


# ---------------------------
# SAVE AS TXT
# ---------------------------
# ensure you have a global variable defined near top:
last_raw_text = ""

def save_as_txt():
    # prefer raw_text for saving (so copy/paste from file is correct)
    text = (last_raw_text or "").strip()
    if not text:
        messagebox.showwarning("Warning", "متنی برای ذخیره وجود ندارد.")
        return

    filename = current_image_name or "ocr_output"
    file_path = filedialog.asksaveasfilename(
        initialfile=f"{filename}.txt",
        defaultextension=".txt",
        filetypes=[("Text File", "*.txt")]
    )
    if not file_path:
        return

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(text)

    messagebox.showinfo("Saved", "فایل متنی ذخیره شد.")


def save_as_docx():
    text = (last_raw_text or "").strip()
    if not text:
        messagebox.showwarning("Warning", "متنی برای ذخیره وجود ندارد.")
        return

    filename = current_image_name or "ocr_output"
    file_path = filedialog.asksaveasfilename(
        initialfile=f"{filename}.docx",
        defaultextension=".docx",
        filetypes=[("Word File", "*.docx")]
    )
    if not file_path:
        return

    doc = Document()
    p = doc.add_paragraph()
    # 2 means RIGHT alignment for python-docx
    p.paragraph_format.alignment = 2
    p.add_run(text)
    doc.save(file_path)

    messagebox.showinfo("Saved", "فایل ورد ذخیره شد.")



def copy_extracted_text():
    try:
        root.clipboard_clear()
        root.clipboard_append(last_raw_text)
        root.update()  # keeps clipboard after closing app
        messagebox.showinfo("کپی شد", "متن در کلیپ‌بورد کپی شد.")
    except Exception as e:
        messagebox.showerror("Error", f"Clipboard error:\n{e}")


# ---------------------------
# TKINTER UI
# ---------------------------
root = tk.Tk()
root.title("Persian OCR - Offline")
root.geometry("900x700")

# Buttons Frame
frame = tk.Frame(root)
frame.pack(pady=10)

# --- Row 0 ---
btn_docx = tk.Button(frame, text="ذخیره DOCX", command=save_as_docx,
                     width=18, height=2, bg="#6A1B9A", fg="white")
btn_docx.grid(row=0, column=0, padx=10)

btn_txt = tk.Button(frame, text="ذخیره TXT", command=save_as_txt,
                    width=18, height=2, bg="#0277BD", fg="white")
btn_txt.grid(row=0, column=1, padx=10)

btn_paste = tk.Button(frame, text="پیست از کلیپ‌بورد", command=paste_from_clipboard,
                      width=20, height=2, bg="#FF9800", fg="white")
btn_paste.grid(row=0, column=2, padx=10)

btn_open = tk.Button(frame, text="انتخاب تصویر", command=run_ocr,
                     width=18, height=2, bg="#4CAF50", fg="white")
btn_open.grid(row=0, column=3, padx=10)

# --- Copy Button (same style, centered below) ---
copy_button = tk.Button(root, text="کپی کردن متن", command=copy_extracted_text,
                        width=20, height=2, bg="#9C27B0", fg="white")
copy_button.pack(pady=8)



# RTL Text Box
text_box = tk.Text(root, width=100, height=30, font=("B Nazanin", 14), wrap="word")
text_box.pack(pady=10)
text_box.tag_configure("right", justify='right')

# Shortcut: Ctrl+V
root.bind("<Control-v>", lambda event: paste_from_clipboard())
root.bind("<Control-V>", lambda event: paste_from_clipboard())

root.mainloop()
